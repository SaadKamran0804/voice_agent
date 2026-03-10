from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_transcript_report(
    summary, conversation_list, output_filename="Call_Transcript.docx"
):
    """
    Generates a formatted Word document (.docx) containing the summary and transcript.
    Takes the summary string and the list of transcript dictionaries.
    """
    # Initialize a new Word document
    document = Document()

    # --- 1. Main Heading ---
    heading = document.add_heading("Transcript", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the participants as a centered, italicized subtitle
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Assistant ↔ User").italic = True

    # --- 2. Summary Section ---
    document.add_heading("Summary", level=1)
    document.add_paragraph(summary)

    # --- 3. Conversation Transcript Section ---
    document.add_heading("Conversation", level=1)

    # Iterate through the transcript list and format the output
    for msg in conversation_list:
        # Grab the role directly from your schema
        role = msg.get("role", "unknown")
        content = msg.get("content", "")

        # Capitalize the role to use as the display name (e.g., "user" -> "User")
        display_name = role.capitalize()

        # Add speaker name in bold
        p_name = document.add_paragraph()
        p_name.add_run(display_name).bold = True

        # Add the actual message content right below it
        document.add_paragraph(content)

    # Save the document
    document.save(output_filename)
    print(f"Document successfully created and saved as: {output_filename}")

    return output_filename
